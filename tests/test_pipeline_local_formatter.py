from __future__ import annotations

from pathlib import Path

from docx import Document

from ai_doc_to_epub.models import ConversionRequest
from ai_doc_to_epub.pipeline import ConversionPipeline


def create_docx(path: Path) -> None:
    doc = Document()
    doc.add_heading("Demo Document", level=1)
    doc.add_paragraph("This is a sample paragraph with a footnote.")
    doc.save(path)


def test_pipeline_with_local_formatter(tmp_path: Path) -> None:
    source = tmp_path / "sample.docx"
    create_docx(source)

    pipeline = ConversionPipeline()
    request = ConversionRequest(
        title="Demo", author="Tester", language="en", use_local_formatter=True
    )

    result = pipeline.convert(source, request)

    assert result.output_path.exists()
    assert result.output_path.suffix == ".epub"
    assert result.file_size > 0
